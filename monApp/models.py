from django.db import models
from django.utils import timezone
from os import listdir
from datetime import datetime, timedelta
from monApp import util
from django.template.loader import get_template
from django.core.mail import EmailMultiAlternatives
from django.utils import timezone
import json
from colour import Color
from docx import Document


# Create your models here.
class ArticleType(models.Model):
    nom = models.CharField(max_length=50, null=True)
    img = models.ImageField(upload_to="static/img/article_icon", null=True)



class ForumGroup(models.Model):

    def __str__(self):
        return self.nom

    nom = models.CharField(max_length=30)
    couleur = models.CharField(max_length=10)
    logo = models.ForeignKey("logo", on_delete=models.SET_NULL, null=True)
    mail = models.EmailField(null=True, max_length=50)

    list_template = "monApp/admin_group_list.html"
    form_template = "monApp/admin_group_form.html"
    form_fields = ("nom", "couleur", "logo")


    @staticmethod
    def objectList(arg):
        return ForumGroup.objects.all()

    def couleur2(self):
        try:
            c2 = Color(self.couleur)
        except ValueError:
            return "#ffffff"
        l = c2.get_luminance()
        l = (l+1)/2
        c2.set_luminance(l)
        return c2.get_hex()



class LangueSection(models.Model):
    nom_complet = models.CharField(max_length=50)
    nom_apparent = models.CharField(max_length=50)
    groupes = models.ManyToManyField(ForumGroup)

    list_template = "monApp/admin_section_list.html"
    form_template = "monApp/admin_section_form.html"

    def __str__(self):
        return self.nom_complet

    def getArticles(self):
        return Article.objects.filter(langue=self, date__lte=timezone.localdate())

    @staticmethod
    def objectList(arg):
        return LangueSection.objects.all()


class ForumUser(models.Model):
    identifiant = models.CharField(max_length=50, unique=True, null=False)
    password = models.CharField(max_length=300, null=False)
    mail = models.EmailField(null = True)
    groupe = models.ManyToManyField(ForumGroup)
    admin = models.BooleanField(default=False)

    list_template = "monApp/admin_user_list.html"
    form_template = "monApp/admin_user_form.html"

    def __str__(self):
        return self.identifiant


    def getLangues(self):
        langues = []
        for groupe in self.groupe.all():
            langues_groupe = LangueSection.objects.filter(groupes=groupe)
            for langue in langues_groupe:
                if not langue in langues:
                    langues.append(langue)
        return langues

    @staticmethod
    def objectList(arg):
        if arg:
            return ForumUser.objects.filter(groupe=arg)
        else:
            return ForumUser.objects.all()

    def send_init_mail(self):
        if self.mail:
            file = open("sent_mails", "w")
            template = get_template("monApp/mail_templates/addUser.txt")
            mail_content = template.render({"username": self.identifiant,
                                            "password": self.password})
            msg = EmailMultiAlternatives("Bienvenue", mail_content, "noreply@docs-ceci-formation.fr", [self.mail])
            msg.send()
            self.password = util.hash(self.password)
            self.save()
            file.write(self.mail)
        else:
            raise Exception("User has no E-mail")

    def send_newpassword(self):
        if self.mail:
            self.password = util.gen_passwd(10)
            template = get_template("monApp/mail_templates/newPassword.txt")
            mail_content = template.render({"username": self.identifiant,
                                            "password":self.password})
            self.password = util.hash(self.password)
            msg = EmailMultiAlternatives('Demande de nouveau mot de passe', mail_content, "noreply@docs-ceci-formation.fr", [self.mail])
            msg.send()
            self.save()
        else:
            raise Exception("User has no e-mail")

class logo(models.Model):
    nom = models.CharField(max_length=50)
    img = models.ImageField(upload_to="static/img/logo")


class Article(models.Model):
    title = models.CharField(max_length=30, null=False)
    subTitle = models.CharField(max_length=100, null=True)
    date = models.DateField(default=timezone.now)
    type_id = models.ForeignKey("ArticleType", on_delete=models.PROTECT, default=1)
    langue = models.ManyToManyField(LangueSection)

    list_template = "monApp/admin_article_list.html"
    form_template = "monApp/admin_article_form.html"

    class Meta:
        verbose_name="article"
        ordering=["date"]

    def __str__(self):
        return self.title

    def get_images_path(self):
        path = "monApp/static/content/%s/"%(self.id)
        try:
            fileList = listdir(path)
            fileList.sort()
            print (fileList)
            return ["content/%s/%s"%(self.id, name) for name in fileList]
        except Exception:
            return []

    def get_audio_path(self):
        path = "monApp/static/content_audio/%s"%(self.id)
        try:
            fileList = listdir(path)
            fileList.sort()
            return ["content_audio/%s/%s"%(self.id, name) for name in fileList]
        except Exception:
            return []

    @staticmethod
    def objectList(arg=None):
        if arg:
            return Article.objects.filter(langue=arg)
        else:
            return Article.objects.all()


class PlannedMail(models.Model):
    to = models.ForeignKey("ForumGroup", on_delete=models.CASCADE)
    subject = models.CharField(null=True, max_length=50)
    content = models.TextField(null=True)
    sent = models.BooleanField(default=False)
    time = models.DateField(default=timezone.now)

    list_template = "monApp/admin_mail_list.html"
    form_template = "monApp/admin_mail_form_test.html"

    @staticmethod
    def objectList(arg=None):
        return PlannedMail.objects.all().order_by("time")

    def __str__(self):
        return "%s - %s - %s/%s/%s" % (self.subject,
                                             self.to,
                                             self.time.day,
                                             self.time.month,
                                             self.time.year,
                                             )

    def send(self):
        for user in self.to.forumuser_set.all():
                txt_content = self.content
                msg = EmailMultiAlternatives(self.subject, txt_content, "noreply@docs-ceci-formation.fr", [user.mail])
                ret = msg.send()
        if ret:
            self.delete()
        else:
            util.report("Erreur d'envoi d'un mail")

class ExamInvite(models.Model):
    to = models.ForeignKey("ForumGroup", on_delete=models.CASCADE)
    date = models.DateTimeField(default=timezone.now)
    examType = models.CharField(default="Toeic", max_length=20)

    list_template = "monApp/admin_exam_list.html"
    form_template = "monApp/admin_exam_form.html"

    def get_nb_accept(self):
        return ExamInviteItem.objects.filter(parent=self.id, accepted=True).count()

    @staticmethod
    def objectList(arg=None):
        return ExamInvite.objects.filter(date__gte=datetime.now()).order_by("date")

    def __str__(self):
        return "%s - %s - %s" % (self.examType, self.to, self.date)

class ExamInviteItem(models.Model):
    user = models.ForeignKey("ForumUser", on_delete=models.CASCADE)
    parent = models.ForeignKey("ExamInvite", on_delete=models.CASCADE)
    accepted = models.BooleanField(default=False)
    link = models.TextField()


    @staticmethod
    def objectList(arg=None):
        return ExamInviteItem.objects.filter(parent=arg)

    def __str__(self):
        return self.user.identifiant

    def send(self):
        template = get_template("monApp/mail_templates/exam_invite.txt")
        text = template.render({"parent": self.parent, "lien": self.link })
        mail = EmailMultiAlternatives("%s du %s"%(self.parent.examType, str(self.parent.date.date())), text, "noreply@docs-ceci-formation.fr", [self.user.mail])
        mail.send()

    def generate_convoc(self):

        doc = Document("monApp/static/toeic.docx")

        doc.paragraphs[0].runs[2].text = self.parent.examType
        doc.paragraphs[1].runs[0].text = self.user.identifiant.replace(".", " ").capitalize()
        date = self.parent.date
        doc.tables[0]._cells[0].paragraphs[1].runs[2].text = date.strftime("%d/%m/%Y")
        doc.tables[0]._cells[0].paragraphs[2].runs[1].text = (date-timedelta(hours=1)).strftime("%Hh%M")
        doc.tables[0]._cells[0].paragraphs[2].runs[6].text = date.strftime("%Hh%M")
        doc.tables[0]._cells[0].paragraphs[2].runs[8].text = "(Durée %s)" % ("1h" if self.parent.examType == "Widaf" else "2h30")


        doc.save("monApp/convocations/%s - %s - %s.docx"%(self.parent.examType, self.user, self.link))
        return("monApp/convocations/%s - %s - %s.docx"%(self.parent.examType, self.user, self.link))

    def send_confirm(self, convoc):
        template = get_template("monApp/mail_templates/exam_confirm.txt")
        text = template.render({"parent": self.parent })
        mail = EmailMultiAlternatives("Confirmation - %s du %s"%(self.parent.examType, str(self.parent.date.date())), text, "noreply@docs-ceci-formation.fr", [self.user.mail])
        mail.attach_file(convoc)
        mail.send()


